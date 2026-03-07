"""
JAWIR OS - File Generator
=========================
Generate various file formats using Python.

Supported formats:
- Word (.docx)
- PDF
- CSV
- Excel (.xlsx)
- JSON
- TXT
- Markdown
- Images/Charts (PNG, JPG)
"""

import os
import sys
import json
import csv
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
from datetime import datetime


class FileGenerator:
    """
    File Generator untuk JAWIR OS.
    Membuat berbagai format file secara programatis.
    """
    
    def __init__(self, output_dir: str = None):
        self.output_dir = Path(output_dir or "D:/sijawir/python_workspace/output").absolute()
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def _get_output_path(self, filename: str, extension: str) -> Path:
        """Generate output path with timestamp if file exists."""
        if not filename.endswith(extension):
            filename = f"{filename}{extension}"
        
        path = self.output_dir / filename
        
        # If file exists, add timestamp
        if path.exists():
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            name = path.stem
            path = self.output_dir / f"{name}_{timestamp}{extension}"
        
        return path
    
    # ==================== WORD DOCUMENT ====================
    
    def create_word(
        self,
        content: Union[str, List[Dict]],
        filename: str = "document",
        title: str = None,
        author: str = "JAWIR OS"
    ) -> Dict[str, Any]:
        """
        Create a Word document (.docx).
        
        Args:
            content: Text content or list of paragraphs with formatting
                    e.g., [{"text": "Hello", "bold": True, "size": 14}]
            filename: Output filename (without extension)
            title: Document title
            author: Document author
            
        Returns:
            Dict with success status and file path
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return {
                "success": False,
                "message": "python-docx not installed. Run: pip install python-docx"
            }
        
        try:
            doc = Document()
            
            # Set document properties
            doc.core_properties.author = author
            if title:
                doc.core_properties.title = title
                # Add title
                title_para = doc.add_heading(title, 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content
            if isinstance(content, str):
                # Simple text content
                for paragraph in content.split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            elif isinstance(content, list):
                # Formatted content
                for item in content:
                    if isinstance(item, str):
                        doc.add_paragraph(item)
                    elif isinstance(item, dict):
                        text = item.get("text", "")
                        style = item.get("style", "Normal")
                        
                        if item.get("heading"):
                            para = doc.add_heading(text, level=item.get("level", 1))
                        else:
                            para = doc.add_paragraph(style=style)
                            run = para.add_run(text)
                            
                            if item.get("bold"):
                                run.bold = True
                            if item.get("italic"):
                                run.italic = True
                            if item.get("underline"):
                                run.underline = True
                            if item.get("size"):
                                run.font.size = Pt(item["size"])
            
            # Save
            output_path = self._get_output_path(filename, ".docx")
            doc.save(str(output_path))
            
            return {
                "success": True,
                "message": f"Word document created: {output_path}",
                "path": str(output_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating Word document: {str(e)}"
            }
    
    # ==================== PDF ====================
    
    def create_pdf(
        self,
        content: Union[str, List[Dict]],
        filename: str = "document",
        title: str = None,
        author: str = "JAWIR OS"
    ) -> Dict[str, Any]:
        """
        Create a PDF document.
        
        Args:
            content: Text content or list of elements
            filename: Output filename (without extension)
            title: Document title
            author: Document author
            
        Returns:
            Dict with success status and file path
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        except ImportError:
            return {
                "success": False,
                "message": "reportlab not installed. Run: pip install reportlab"
            }
        
        try:
            output_path = self._get_output_path(filename, ".pdf")
            
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72,
                author=author,
                title=title or filename
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            if title:
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    alignment=TA_CENTER,
                    fontSize=24,
                    spaceAfter=30
                )
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 12))
            
            # Add content
            if isinstance(content, str):
                for paragraph in content.split('\n\n'):
                    if paragraph.strip():
                        story.append(Paragraph(paragraph.strip(), styles['Normal']))
                        story.append(Spacer(1, 12))
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, str):
                        story.append(Paragraph(item, styles['Normal']))
                        story.append(Spacer(1, 6))
                    elif isinstance(item, dict):
                        text = item.get("text", "")
                        style_name = item.get("style", "Normal")
                        
                        if item.get("heading"):
                            level = item.get("level", 1)
                            style_name = f"Heading{level}"
                        
                        if style_name in styles:
                            story.append(Paragraph(text, styles[style_name]))
                        else:
                            story.append(Paragraph(text, styles['Normal']))
                        
                        story.append(Spacer(1, 6))
            
            doc.build(story)
            
            return {
                "success": True,
                "message": f"PDF created: {output_path}",
                "path": str(output_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating PDF: {str(e)}"
            }
    
    # ==================== CSV ====================
    
    def create_csv(
        self,
        data: List[Dict],
        filename: str = "data",
        headers: List[str] = None
    ) -> Dict[str, Any]:
        """
        Create a CSV file.
        
        Args:
            data: List of dictionaries (each dict is a row)
            filename: Output filename (without extension)
            headers: Optional column headers (auto-detected if not provided)
            
        Returns:
            Dict with success status and file path
        """
        try:
            output_path = self._get_output_path(filename, ".csv")
            
            if not data:
                return {
                    "success": False,
                    "message": "No data provided"
                }
            
            # Auto-detect headers
            if not headers:
                headers = list(data[0].keys())
            
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=headers)
                writer.writeheader()
                writer.writerows(data)
            
            return {
                "success": True,
                "message": f"CSV created: {output_path}",
                "path": str(output_path),
                "rows": len(data),
                "columns": len(headers)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating CSV: {str(e)}"
            }
    
    # ==================== EXCEL ====================
    
    def create_excel(
        self,
        data: Union[List[Dict], Dict[str, List[Dict]]],
        filename: str = "workbook",
        sheet_name: str = "Sheet1"
    ) -> Dict[str, Any]:
        """
        Create an Excel file (.xlsx).
        
        Args:
            data: List of dicts (single sheet) or Dict of sheet_name -> List of dicts (multiple sheets)
            filename: Output filename (without extension)
            sheet_name: Sheet name for single sheet mode
            
        Returns:
            Dict with success status and file path
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            from openpyxl.utils.dataframe import dataframe_to_rows
        except ImportError:
            return {
                "success": False,
                "message": "openpyxl not installed. Run: pip install openpyxl"
            }
        
        try:
            wb = Workbook()
            
            # Header style
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            
            def write_sheet(ws, rows_data, name=None):
                if name:
                    ws.title = name
                
                if not rows_data:
                    return
                
                # Write headers
                headers = list(rows_data[0].keys())
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                
                # Write data
                for row_idx, row_data in enumerate(rows_data, 2):
                    for col_idx, header in enumerate(headers, 1):
                        ws.cell(row=row_idx, column=col_idx, value=row_data.get(header, ""))
                
                # Auto-adjust column width
                for col in ws.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column].width = adjusted_width
            
            # Handle single sheet or multiple sheets
            if isinstance(data, list):
                ws = wb.active
                write_sheet(ws, data, sheet_name)
            elif isinstance(data, dict):
                first = True
                for sheet_name, sheet_data in data.items():
                    if first:
                        ws = wb.active
                        first = False
                    else:
                        ws = wb.create_sheet()
                    write_sheet(ws, sheet_data, sheet_name)
            
            output_path = self._get_output_path(filename, ".xlsx")
            wb.save(str(output_path))
            
            return {
                "success": True,
                "message": f"Excel created: {output_path}",
                "path": str(output_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating Excel: {str(e)}"
            }
    
    # ==================== JSON ====================
    
    def create_json(
        self,
        data: Any,
        filename: str = "data",
        pretty: bool = True
    ) -> Dict[str, Any]:
        """
        Create a JSON file.
        
        Args:
            data: Any JSON-serializable data
            filename: Output filename (without extension)
            pretty: Pretty print with indentation
            
        Returns:
            Dict with success status and file path
        """
        try:
            output_path = self._get_output_path(filename, ".json")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                if pretty:
                    json.dump(data, f, indent=2, ensure_ascii=False)
                else:
                    json.dump(data, f, ensure_ascii=False)
            
            return {
                "success": True,
                "message": f"JSON created: {output_path}",
                "path": str(output_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating JSON: {str(e)}"
            }
    
    # ==================== TXT ====================
    
    def create_txt(
        self,
        content: str,
        filename: str = "document"
    ) -> Dict[str, Any]:
        """
        Create a text file.
        
        Args:
            content: Text content
            filename: Output filename (without extension)
            
        Returns:
            Dict with success status and file path
        """
        try:
            output_path = self._get_output_path(filename, ".txt")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                "success": True,
                "message": f"TXT created: {output_path}",
                "path": str(output_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating TXT: {str(e)}"
            }
    
    # ==================== MARKDOWN ====================
    
    def create_markdown(
        self,
        content: Union[str, List[Dict]],
        filename: str = "document",
        title: str = None
    ) -> Dict[str, Any]:
        """
        Create a Markdown file.
        
        Args:
            content: Markdown content string or structured content
            filename: Output filename (without extension)
            title: Document title (will be added as H1)
            
        Returns:
            Dict with success status and file path
        """
        try:
            output_path = self._get_output_path(filename, ".md")
            
            md_content = ""
            
            # Add title
            if title:
                md_content += f"# {title}\n\n"
            
            # Process content
            if isinstance(content, str):
                md_content += content
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, str):
                        md_content += f"{item}\n\n"
                    elif isinstance(item, dict):
                        text = item.get("text", "")
                        
                        if item.get("heading"):
                            level = item.get("level", 1)
                            md_content += f"{'#' * level} {text}\n\n"
                        elif item.get("bullet"):
                            md_content += f"- {text}\n"
                        elif item.get("numbered"):
                            md_content += f"1. {text}\n"
                        elif item.get("code"):
                            lang = item.get("lang", "")
                            md_content += f"```{lang}\n{text}\n```\n\n"
                        elif item.get("quote"):
                            md_content += f"> {text}\n\n"
                        elif item.get("bold"):
                            md_content += f"**{text}**\n\n"
                        elif item.get("italic"):
                            md_content += f"*{text}*\n\n"
                        else:
                            md_content += f"{text}\n\n"
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            return {
                "success": True,
                "message": f"Markdown created: {output_path}",
                "path": str(output_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating Markdown: {str(e)}"
            }
    
    # ==================== IMAGES & CHARTS ====================
    
    def create_chart(
        self,
        data: Dict[str, List],
        chart_type: str = "line",
        filename: str = "chart",
        title: str = None,
        xlabel: str = None,
        ylabel: str = None,
        figsize: tuple = (10, 6)
    ) -> Dict[str, Any]:
        """
        Create a chart/graph as image.
        
        Args:
            data: Dict with "x" and "y" keys, or multiple series {"series1": [...], "series2": [...]}
            chart_type: "line", "bar", "pie", "scatter", "histogram"
            filename: Output filename (without extension)
            title: Chart title
            xlabel: X-axis label
            ylabel: Y-axis label
            figsize: Figure size (width, height) in inches
            
        Returns:
            Dict with success status and file path
        """
        try:
            import matplotlib.pyplot as plt
            import matplotlib
            matplotlib.use('Agg')  # Non-interactive backend
        except ImportError:
            return {
                "success": False,
                "message": "matplotlib not installed. Run: pip install matplotlib"
            }
        
        try:
            fig, ax = plt.subplots(figsize=figsize)
            
            if chart_type == "pie":
                # Pie chart
                labels = data.get("labels", list(data.keys()))
                values = data.get("values", list(data.values()))
                if isinstance(values[0], list):
                    values = [sum(v) if isinstance(v, list) else v for v in values]
                ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
                ax.axis('equal')
                
            elif chart_type == "histogram":
                # Histogram
                values = data.get("values", data.get("y", []))
                bins = data.get("bins", 10)
                ax.hist(values, bins=bins, edgecolor='black')
                
            else:
                # Line, bar, scatter
                x = data.get("x")
                
                if x is None:
                    # Multiple series without explicit x
                    for label, values in data.items():
                        if label not in ["x", "labels", "values"]:
                            if chart_type == "line":
                                ax.plot(values, label=label, marker='o')
                            elif chart_type == "bar":
                                x_pos = range(len(values))
                                ax.bar(x_pos, values, label=label)
                            elif chart_type == "scatter":
                                ax.scatter(range(len(values)), values, label=label)
                else:
                    y = data.get("y", [])
                    if chart_type == "line":
                        ax.plot(x, y, marker='o')
                    elif chart_type == "bar":
                        ax.bar(x, y)
                    elif chart_type == "scatter":
                        ax.scatter(x, y)
                
                if len(data) > 2:
                    ax.legend()
            
            if title:
                ax.set_title(title, fontsize=14, fontweight='bold')
            if xlabel:
                ax.set_xlabel(xlabel)
            if ylabel:
                ax.set_ylabel(ylabel)
            
            ax.grid(True, alpha=0.3)
            
            output_path = self._get_output_path(filename, ".png")
            plt.savefig(str(output_path), dpi=150, bbox_inches='tight')
            plt.close()
            
            return {
                "success": True,
                "message": f"Chart created: {output_path}",
                "path": str(output_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating chart: {str(e)}"
            }
    
    def create_image(
        self,
        width: int = 800,
        height: int = 600,
        color: str = "white",
        filename: str = "image",
        text: str = None,
        text_color: str = "black"
    ) -> Dict[str, Any]:
        """
        Create a simple image with optional text.
        
        Args:
            width: Image width in pixels
            height: Image height in pixels
            color: Background color
            filename: Output filename (without extension)
            text: Optional text to add
            text_color: Text color
            
        Returns:
            Dict with success status and file path
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
        except ImportError:
            return {
                "success": False,
                "message": "Pillow not installed. Run: pip install Pillow"
            }
        
        try:
            img = Image.new('RGB', (width, height), color=color)
            
            if text:
                draw = ImageDraw.Draw(img)
                # Try to use a nice font, fallback to default
                try:
                    font = ImageFont.truetype("arial.ttf", 32)
                except:
                    font = ImageFont.load_default()
                
                # Center text
                bbox = draw.textbbox((0, 0), text, font=font)
                text_width = bbox[2] - bbox[0]
                text_height = bbox[3] - bbox[1]
                x = (width - text_width) // 2
                y = (height - text_height) // 2
                
                draw.text((x, y), text, fill=text_color, font=font)
            
            output_path = self._get_output_path(filename, ".png")
            img.save(str(output_path))
            
            return {
                "success": True,
                "message": f"Image created: {output_path}",
                "path": str(output_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Error creating image: {str(e)}"
            }
