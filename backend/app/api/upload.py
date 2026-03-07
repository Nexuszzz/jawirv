"""
JAWIR OS - File Upload Handler
Handles image and PDF file uploads with processing.
"""

import base64
import logging
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel

logger = logging.getLogger("jawir.upload")

# Create router
router = APIRouter(prefix="/api/upload", tags=["upload"])

# Upload settings
UPLOAD_DIR = Path("uploads")
ALLOWED_IMAGE_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp"}
ALLOWED_PDF_TYPES = {"application/pdf"}
ALLOWED_DOC_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword",  # .doc
    "text/plain",  # .txt
    "text/markdown",  # .md
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
    "application/vnd.ms-excel",  # .xls
    "text/csv",  # .csv
    "application/json",  # .json
}
ALLOWED_TYPES = ALLOWED_IMAGE_TYPES | ALLOWED_PDF_TYPES | ALLOWED_DOC_TYPES
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB


class UploadResponse(BaseModel):
    """Response model for file upload."""
    success: bool
    file_id: str
    filename: str
    file_type: str
    file_size: int
    url: str
    base64_preview: Optional[str] = None
    message: str


class FileInfo(BaseModel):
    """File information model."""
    file_id: str
    filename: str
    file_type: str
    file_size: int
    upload_time: str
    url: str


# Ensure upload directory exists
UPLOAD_DIR.mkdir(exist_ok=True)


def get_file_extension(content_type: str) -> str:
    """Get file extension from content type."""
    extensions = {
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/gif": ".gif",
        "image/webp": ".webp",
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",
        "text/plain": ".txt",
        "text/markdown": ".md",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
        "application/vnd.ms-excel": ".xls",
        "text/csv": ".csv",
        "application/json": ".json",
    }
    return extensions.get(content_type, ".bin")


def generate_file_id() -> str:
    """Generate unique file ID."""
    return f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"


@router.post("/file", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    Upload a file (image or PDF).
    
    Supports:
    - Images: JPEG, PNG, GIF, WebP
    - Documents: PDF
    
    Max size: 10MB
    """
    # Validate content type
    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{file.content_type}' not allowed. "
                   f"Allowed: {', '.join(ALLOWED_TYPES)}"
        )
    
    # Read file content
    content = await file.read()
    file_size = len(content)
    
    # Validate file size
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Max size: {MAX_FILE_SIZE // (1024*1024)}MB"
        )
    
    # Generate file ID and path
    file_id = generate_file_id()
    extension = get_file_extension(file.content_type)
    safe_filename = f"{file_id}{extension}"
    file_path = UPLOAD_DIR / safe_filename
    
    # Save file
    try:
        with open(file_path, "wb") as f:
            f.write(content)
        logger.info(f"📁 File uploaded: {safe_filename} ({file_size} bytes)")
    except Exception as e:
        logger.error(f"Failed to save file: {e}")
        raise HTTPException(status_code=500, detail="Failed to save file")
    
    # Generate base64 preview for images
    base64_preview = None
    if file.content_type in ALLOWED_IMAGE_TYPES:
        base64_preview = base64.b64encode(content).decode("utf-8")
    
    original_name = file.filename or safe_filename
    
    return UploadResponse(
        success=True,
        file_id=file_id,
        filename=original_name,
        file_type=file.content_type,
        file_size=file_size,
        url=f"/api/upload/files/{safe_filename}",
        base64_preview=base64_preview,
        message=f"File '{original_name}' berhasil diupload!"
    )


@router.get("/files/{filename}")
async def get_file(filename: str):
    """Serve uploaded file."""
    from fastapi.responses import FileResponse
    
    file_path = UPLOAD_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    # Determine media type
    media_type = "application/octet-stream"
    if filename.endswith(".jpg") or filename.endswith(".jpeg"):
        media_type = "image/jpeg"
    elif filename.endswith(".png"):
        media_type = "image/png"
    elif filename.endswith(".gif"):
        media_type = "image/gif"
    elif filename.endswith(".webp"):
        media_type = "image/webp"
    elif filename.endswith(".pdf"):
        media_type = "application/pdf"
    
    return FileResponse(file_path, media_type=media_type)


@router.delete("/files/{file_id}")
async def delete_file(file_id: str):
    """Delete an uploaded file."""
    # Find file with matching ID
    for file_path in UPLOAD_DIR.iterdir():
        if file_path.stem.startswith(file_id):
            try:
                file_path.unlink()
                logger.info(f"🗑️ File deleted: {file_path.name}")
                return {"success": True, "message": "File berhasil dihapus"}
            except Exception as e:
                logger.error(f"Failed to delete file: {e}")
                raise HTTPException(status_code=500, detail="Failed to delete file")
    
    raise HTTPException(status_code=404, detail="File not found")


@router.get("/list")
async def list_files():
    """List all uploaded files."""
    files = []
    for file_path in UPLOAD_DIR.iterdir():
        if file_path.is_file():
            stat = file_path.stat()
            files.append({
                "filename": file_path.name,
                "file_size": stat.st_size,
                "upload_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "url": f"/api/upload/files/{file_path.name}",
            })
    
    return {
        "files": sorted(files, key=lambda x: x["upload_time"], reverse=True),
        "total": len(files),
    }


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from PDF file.
    Requires PyMuPDF (fitz) or pdfplumber.
    """
    try:
        import fitz  # PyMuPDF
        
        text_parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text())
        
        return "\n\n".join(text_parts)
    
    except ImportError:
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            logger.warning("No PDF library available (PyMuPDF or pdfplumber)")
            return "[PDF text extraction not available]"


def prepare_image_for_gemini(file_path: Path) -> dict:
    """
    Prepare image for Gemini API multimodal input.
    
    Returns:
        Dict with mime_type and base64 data
    """
    with open(file_path, "rb") as f:
        content = f.read()
    
    # Determine MIME type
    suffix = file_path.suffix.lower()
    mime_types = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }
    mime_type = mime_types.get(suffix, "application/octet-stream")
    
    return {
        "mime_type": mime_type,
        "data": base64.b64encode(content).decode("utf-8"),
    }


def get_file_for_agent(file_id: str) -> Optional[dict]:
    """
    Get file info and content for agent processing.
    
    Returns:
        Dict with file info and content (text for PDF/docs, image data for images)
    """
    for file_path in UPLOAD_DIR.iterdir():
        if file_path.stem.startswith(file_id):
            suffix = file_path.suffix.lower()
            
            if suffix == ".pdf":
                # Extract text from PDF
                text = extract_text_from_pdf(file_path)
                return {
                    "type": "pdf",
                    "filename": file_path.name,
                    "content": text,
                }
            
            elif suffix == ".docx":
                # Extract text from Word document
                text = extract_text_from_docx(file_path)
                return {
                    "type": "docx",
                    "filename": file_path.name,
                    "content": text,
                }
            
            elif suffix in [".txt", ".md", ".csv", ".json"]:
                # Read text files directly
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text = f.read()
                    return {
                        "type": "text",
                        "filename": file_path.name,
                        "content": text,
                    }
                except Exception as e:
                    logger.error(f"Failed to read text file: {e}")
                    return {
                        "type": "text",
                        "filename": file_path.name,
                        "content": f"[Error membaca file: {str(e)}]",
                    }
            
            elif suffix in [".xlsx", ".xls"]:
                # Extract data from Excel
                text = extract_text_from_excel(file_path)
                return {
                    "type": "excel",
                    "filename": file_path.name,
                    "content": text,
                }
            
            elif suffix in [".jpg", ".jpeg", ".png", ".gif", ".webp"]:
                # Prepare image for Gemini
                image_data = prepare_image_for_gemini(file_path)
                return {
                    "type": "image",
                    "filename": file_path.name,
                    "mime_type": image_data["mime_type"],
                    "data": image_data["data"],
                }
    
    return None


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from Word document (.docx).
    Requires python-docx.
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    except ImportError:
        logger.warning("python-docx not installed, cannot extract Word content")
        return "[Word text extraction not available - install python-docx]"
    except Exception as e:
        logger.error(f"Error extracting Word content: {e}")
        return f"[Error membaca Word: {str(e)}]"


def extract_text_from_excel(file_path: Path) -> str:
    """
    Extract text from Excel file (.xlsx/.xls).
    Requires openpyxl or xlrd.
    """
    try:
        import openpyxl
        
        wb = openpyxl.load_workbook(file_path, data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip() and row_text.replace("|", "").strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    except ImportError:
        logger.warning("openpyxl not installed, cannot extract Excel content")
        return "[Excel extraction not available - install openpyxl]"
    except Exception as e:
        logger.error(f"Error extracting Excel content: {e}")
        return f"[Error membaca Excel: {str(e)}]"
